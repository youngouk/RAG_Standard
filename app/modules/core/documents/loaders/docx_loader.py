"""
DOCX Document Loader
Word 문서 파일 로딩 전략 구현
"""

from pathlib import Path

from docx import Document as DocxDocument
from langchain_core.documents import Document

from .....lib.logger import get_logger
from .base import DocumentLoaderStrategy

logger = get_logger(__name__)


class DOCXLoader(DocumentLoaderStrategy):
    """Word 문서 로더"""

    @property
    def supported_extensions(self) -> list[str]:
        return [".docx", ".DOCX", ".doc", ".DOC"]

    async def load(self, file_path: Path) -> list[Document]:
        """DOCX 파일 로드"""
        try:
            doc = DocxDocument(file_path)
            paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
            if not paragraphs:
                logger.warning(f"Empty DOCX file: {file_path.name}")
                return []
            content = "\n".join(paragraphs)
            documents = [Document(page_content=content, metadata={})]
            logger.info(f"DOCX loaded: {len(paragraphs)} paragraphs from {file_path.name}")
            return documents
        except Exception as e:
            logger.error(f"DOCX loading failed for {file_path}: {e}")
            raise ValueError(f"Failed to load DOCX file: {e}") from e
